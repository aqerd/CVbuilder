import os

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT

from settings import Settings

regular_font = "Arial"
bold_font = "Arial-Bold"

regular_font_path = Settings.REGULAR_FONT_PATH
bold_font_path = Settings.BOLD_FONT_PATH

pdfmetrics.registerFont(TTFont(regular_font, regular_font_path))
pdfmetrics.registerFont(TTFont(bold_font, bold_font_path))

PATH_SAVE = os.path.join(os.path.dirname(__file__), "files")

if not os.path.exists(PATH_SAVE):
    os.makedirs(PATH_SAVE)


def create_type(data, file_type):
    match file_type:
        case "pdf":
            filename = create_pdf(data)
        case "doc":
            filename = create_docx(data)
        case "jpg":
            filename = create_jpg(data)
        case _:
            return "File type not found", 404
    return filename


def create_pdf(data):
    pdf_path = os.path.join(PATH_SAVE, "CV.pdf")
    styles = getSampleStyleSheet()
    styleN = styles["BodyText"]
    styleN.alignment = TA_LEFT
    styleH = styles["Heading1"]
    styleH.alignment = TA_LEFT

    elements = []

    elements.append(Paragraph(f"{data['name']} {data['middle_name']} {data['last_name']}", styleH))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph(f"Age: {data['age']}", styleN))
    elements.append(Paragraph(f"Date of Birth: {data['dob']}", styleN))
    elements.append(Paragraph(f"Email: {data['email']}", styleN))
    elements.append(Paragraph(f"Country: {data['country']}", styleN))
    elements.append(Paragraph(f"City: {data['city']}", styleN))
    elements.append(Spacer(1, 12))

    if data.get("socials"):
        elements.append(Paragraph("Socials", styles["h2"]))
        for social in data["socials"]:
            link = f'<a href="{social["link"]}" color="blue">{social["service"]}</a>'
            elements.append(Paragraph(link, styleN))
        elements.append(Spacer(1, 12))

    if data.get("projects"):
        elements.append(Paragraph("Projects", styles["h2"]))
        for project in data["projects"]:
            elements.append(Paragraph(f"Project Name: {project['name']}", styleN))
            elements.append(Paragraph(f"Time: {project['time']}", styleN))
            if project.get("link"):
                link = f'<a href="{project["link"]}" color="blue">{project["link"]}</a>'
                elements.append(Paragraph(link, styleN))
            elements.append(Paragraph(project['description'], styleN))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 12))

    if data.get("experiences"):
        elements.append(Paragraph("Experience", styles["h2"]))
        for exp in data["experiences"]:
            elements.append(Paragraph(f"Company: {exp['company']}", styleN))
            elements.append(Paragraph(f"Job Title: {exp['title']}", styleN))
            elements.append(Paragraph(f"Period: {exp['period']}", styleN))
            elements.append(Paragraph(exp['description'], styleN))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 12))

    if data.get("education"):
        elements.append(Paragraph("Education", styles["h2"]))
        for edu in data["education"]:
            elements.append(Paragraph(f"Institution: {edu['institution']}", styleN))
            elements.append(Paragraph(f"Period: {edu['period']}", styleN))
            elements.append(Paragraph(f"Field of Study: {edu['field']}", styleN))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 12))

    if data.get("languages"):
        elements.append(Paragraph("Languages", styles["h2"]))
        for lang in data["languages"]:
            elements.append(Paragraph(f"{lang['language']}: {lang['level']}", styleN))
            elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 12))

    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    doc.build(elements)

    return pdf_path


def create_docx(data):
    docx_path = os.path.join(PATH_SAVE, "CV.docx")
    doc = Document()
    doc.add_heading(f"{data['name']} {data['middle_name']} {data['last_name']}", 0)
    doc.add_paragraph(f"Age: {data['age']}")
    doc.add_paragraph(f"Date of Birth: {data['dob']}")
    doc.add_paragraph(f"Email: {data['email']}")
    doc.add_paragraph(f"Country: {data['country']}")
    doc.add_paragraph(f"City: {data['city']}")

    if data["socials"]:
        doc.add_heading("Socials", level=1)
        for social in data["socials"]:
            doc.add_paragraph(f"{social['service']}: {social['link']}")

    if data["projects"]:
        doc.add_heading("Projects", level=1)
        for project in data["projects"]:
            doc.add_paragraph(f"Project Name: {project['name']}")
            doc.add_paragraph(f"Time: {project['time']}")
            doc.add_paragraph(f"Link: {project['link']}")
            doc.add_paragraph(project['description'])

    if data["experiences"]:
        doc.add_heading("Experience", level=1)
        for exp in data["experiences"]:
            doc.add_paragraph(f"Company: {exp['company']}")
            doc.add_paragraph(f"Job Title: {exp['title']}")
            doc.add_paragraph(f"Period: {exp['period']}")
            doc.add_paragraph(exp['description'])

    if data["education"]:
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            doc.add_paragraph(f"Institution: {edu['institution']}")
            doc.add_paragraph(f"Period: {edu['period']}")
            doc.add_paragraph(f"Field of Study: {edu['field']}")

    if data["languages"]:
        doc.add_heading("Languages", level=1)
        for lang in data["languages"]:
            doc.add_paragraph(f"{lang['language']}: {lang['level']}")

    doc.save(docx_path)
    return docx_path


def create_jpg(data):
    jpg_path = os.path.join(PATH_SAVE, "CV.jpg")

    width, height = 600, 800
    image = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(image)

    try:
        font_regular = ImageFont.truetype(regular_font_path, 17)
        font_bold = ImageFont.truetype(bold_font_path, 20)
        font_bold_big_size = ImageFont.truetype(bold_font_path, 30)
    except IOError:
        return

    y_position = 30

    draw.text(
        (70, y_position),
        f"{data['name']} {data['middle_name']} {data['last_name']}",
        font=font_bold_big_size,
        fill=(0, 0, 0),
    )
    y_position += 40
    draw.text(
        (100, y_position), f"Age: {data['age']}", font=font_regular, fill=(0, 0, 0)
    )
    y_position += 23
    draw.text(
        (100, y_position),
        f"Date of Birth: {data['dob']}",
        font=font_regular,
        fill=(0, 0, 0),
    )
    y_position += 23
    draw.text(
        (100, y_position), f"Email: {data['email']}", font=font_regular, fill=(0, 0, 0)
    )
    y_position += 23
    draw.text(
        (100, y_position),
        f"Country: {data['country']}",
        font=font_regular,
        fill=(0, 0, 0),
    )
    y_position += 23
    draw.text(
        (100, y_position), f"City: {data['city']}", font=font_regular, fill=(0, 0, 0)
    )
    y_position += 30

    if data["socials"]:
        draw.text((80, y_position), "Socials:", font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for social in data["socials"]:
            draw.text(
                (100, y_position),
                f"{social['service']}: {social['link']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23

    if data["projects"]:
        draw.text((80, y_position), "Projects:", font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for project in data["projects"]:
            draw.text(
                (100, y_position),
                f"Project Name: {project['name']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Time: {project['time']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Link: {project['link']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                project['description'],
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23

    if data["experiences"]:
        draw.text((80, y_position), "Experience:", font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for exp in data["experiences"]:
            draw.text(
                (100, y_position),
                f"Company: {exp['company']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Job Title: {exp['title']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Period: {exp['period']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                exp['description'],
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23

    if data["education"]:
        draw.text((80, y_position), "Education:", font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for edu in data["education"]:
            draw.text(
                (100, y_position),
                f"Institution: {edu['institution']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Period: {edu['period']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23
            draw.text(
                (100, y_position),
                f"Field of Study: {edu['field']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23

    if data["languages"]:
        draw.text((80, y_position), "Languages:", font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for lang in data["languages"]:
            draw.text(
                (100, y_position),
                f"{lang['language']}: {lang['level']}",
                font=font_regular,
                fill=(0, 0, 0),
            )
            y_position += 23

    image.save(jpg_path)
    return jpg_path
