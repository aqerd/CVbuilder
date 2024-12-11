import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from PIL import Image, ImageDraw, ImageFont

regular_font = "Arial"
bold_font = "Arial-Bold"
regular_font_path = "static/fonts/arial.ttf"
bold_font_path = "static/fonts/arialbd.ttf"

pdfmetrics.registerFont(TTFont(regular_font, regular_font_path))
pdfmetrics.registerFont(TTFont(bold_font, bold_font_path))

PATH_SAVE = os.path.join(os.path.dirname(__file__), 'files')

if not os.path.exists(PATH_SAVE):
    os.makedirs(PATH_SAVE)

def collect_data(request):
    name = request.form.get('name')
    middle_name = request.form.get('middle-name')
    last_name = request.form.get('last-name')
    age = request.form.get('age')
    email = request.form.get('email')
    dob = request.form.get('dob')
    citizenship = request.form.get('citizenship')
    city = request.form.get('city')

    socials = []
    social_count = 1
    while request.form.get(f'social-service-{social_count}'):
        social_service = request.form.get(f'social-service-{social_count}')
        social_link = request.form.get(f'social-link-{social_count}')
        socials.append({
            'service': social_service,
            'link': social_link
        })
        social_count += 1

    projects = []
    project_count = 1
    while request.form.get(f'project-name-{project_count}'):
        project_name = request.form.get(f'project-name-{project_count}')
        project_time = request.form.get(f'project-time-{project_count}')
        project_link = request.form.get(f'project-link-{project_count}')
        project_description = request.form.get(f'project-description-{project_count}')

        projects.append({
            'name': project_name,
            'time': project_time,
            'link': project_link,
            'description': project_description
        })

        project_count += 1

    experiences = []
    experience_count = 1
    while request.form.get(f'company-name-{experience_count}'):
        company_name = request.form.get(f'company-name-{experience_count}')
        job_title = request.form.get(f'job-title-{experience_count}')
        work_period = request.form.get(f'work-period-{experience_count}')
        job_description = request.form.get(f'job-description-{experience_count}')

        experiences.append({
            'company': company_name,
            'title': job_title,
            'period': work_period,
            'description': job_description
        })
        experience_count += 1

    education = []
    education_count = 1
    while request.form.get(f'institution-name-{education_count}'):
        institution_name = request.form.get(f'institution-name-{education_count}')
        education_period = request.form.get(f'education-period-{education_count}')
        field_of_study = request.form.get(f'field-of-study-{education_count}')

        education.append({
            'institution': institution_name,
            'period': education_period,
            'field': field_of_study
        })
        education_count += 1

    languages = []
    language_count = 1
    while request.form.get(f'language-name-{language_count}'):
        language_name = request.form.get(f'language-name-{language_count}')
        language_level = request.form.get(f'language-level-{language_count}')
        languages.append({
            'language': language_name,
            'level': language_level
        })
        language_count += 1

    return { 'name': name, 'middle_name': middle_name, 'last_name': last_name, 'age': age, 'email': email,
    'dob': dob, 'citizenship': citizenship, 'city': city, 'socials': socials, 'projects': projects,
    'experiences': experiences, 'education': education, 'languages': languages}

def create_type(data, file_type):
    match file_type:
        case 'pdf':
            filename = create_pdf(data)
        case 'doc':
            filename = create_docx(data)
        case 'jpg':
            filename = create_jpg(data)
        case _:
            return "File type not found", 404
    return filename

def create_pdf(data):
    pdf_path = os.path.join(PATH_SAVE, 'CV.pdf')
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter  
    
    c.setFont(bold_font, 20)
    c.drawString(80, height - 100, f'{data["name"]} {data["middle_name"]} {data["last_name"]}')
    
    c.setFont(regular_font, 12)
    c.drawString(100, height - 130, f'Age: {data["age"]}')
    c.drawString(100, height - 150, f'Date of Birth: {data["dob"]}')
    c.drawString(100, height - 170, f'Email: {data["email"]}')
    c.drawString(100, height - 190, f'Citizenship: {data["citizenship"]}')
    c.drawString(100, height - 210, f'City: {data["city"]}')
    
    y_position = height - 240
    if data["socials"]:
        c.setFont(bold_font, 14)
        c.drawString(80, y_position, 'Social Networks:')  
        y_position -= 20
        c.setFont(regular_font, 12)
        for social in data["socials"]:
            c.drawString(100, y_position, f'{social["service"]}: {social["link"]}')
            y_position -= 20

    if data["projects"]:
        c.setFont(bold_font, 14)
        c.drawString(80, y_position, 'Projects:')  
        y_position -= 20
        c.setFont(regular_font, 12)
        for project in data["projects"]:
            c.drawString(100, y_position, f'Project Name: {project["name"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Time: {project["time"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Link: {project["link"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Description: {project["description"]}')
            y_position -= 20
    
    if data["experiences"]:
        c.setFont(bold_font, 14)
        c.drawString(80, y_position, 'Experience:')  
        y_position -= 20
        c.setFont(regular_font, 12)
        for exp in data["experiences"]:
            c.drawString(100, y_position, f'Company: {exp["company"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Job Title: {exp["title"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Period: {exp["period"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Description: {exp["description"]}')
            y_position -= 20
    
    if data["education"]:
        c.setFont(bold_font, 14)
        c.drawString(80, y_position, 'Education:')  
        y_position -= 20
        c.setFont(regular_font, 12)
        for edu in data["education"]:
            c.drawString(100, y_position, f'Institution: {edu["institution"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Period: {edu["period"]}')
            y_position -= 20
            c.drawString(100, y_position, f'Field of Study: {edu["field"]}')
            y_position -= 20
    
    if data["languages"]:
        c.setFont(bold_font, 14)
        c.drawString(80, y_position, 'Languages:')  
        y_position -= 20
        c.setFont(regular_font, 12)
        for lang in data["languages"]:
            c.drawString(100, y_position, f'{lang["language"]}: {lang["level"]}')
            y_position -= 20
    c.save()

    return pdf_path

def create_docx(data):
    docx_path = os.path.join(PATH_SAVE, 'CV.docx')
    doc = Document()
    doc.add_heading(f'{data["name"]} {data["middle_name"]} {data["last_name"]}', 0)
    doc.add_paragraph(f'Age: {data["age"]}')
    doc.add_paragraph(f'Date of Birth: {data["dob"]}')
    doc.add_paragraph(f'Email: {data["email"]}')
    doc.add_paragraph(f'Citizenship: {data["citizenship"]}')
    doc.add_paragraph(f'City: {data["city"]}')

    if data["socials"]:
        doc.add_heading('Social Networks', level=1)
        for social in data["socials"]:
            doc.add_paragraph(f'{social["service"]}: {social["link"]}')
    
    if data["projects"]:
        doc.add_heading('Projects', level=1)
        for project in data["projects"]:
            doc.add_paragraph(f'Project Name: {project["name"]}')
            doc.add_paragraph(f'Time: {project["time"]}')
            doc.add_paragraph(f'Link: {project["link"]}')
            doc.add_paragraph(f'Description: {project["description"]}')
    
    if data["experiences"]:
        doc.add_heading('Experience', level=1)
        for exp in data["experiences"]:
            doc.add_paragraph(f'Company: {exp["company"]}')
            doc.add_paragraph(f'Job Title: {exp["title"]}')
            doc.add_paragraph(f'Period: {exp["period"]}')
            doc.add_paragraph(f'Description: {exp["description"]}')
    
    if data["education"]:
        doc.add_heading('Education', level=1)
        for edu in data["education"]:
            doc.add_paragraph(f'Institution: {edu["institution"]}')
            doc.add_paragraph(f'Period: {edu["period"]}')
            doc.add_paragraph(f'Field of Study: {edu["field"]}')

    if data["languages"]:
        doc.add_heading('Languages', level=1)
        for lang in data["languages"]:
            doc.add_paragraph(f'{lang["language"]}: {lang["level"]}')
          
    doc.save(docx_path) 
    return docx_path

def create_jpg(data):
    jpg_path = os.path.join(PATH_SAVE, 'CV.jpg')
    
    width, height = 600, 800
    image = Image.new("RGB", (width, height), (255, 255, 255))  
    draw = ImageDraw.Draw(image)

    try:
        font_regular = ImageFont.truetype(regular_font_path, 17)
        font_bold = ImageFont.truetype(bold_font_path, 20)
        font_bold_big_size = ImageFont.truetype(bold_font_path,30 )
    except IOError:
        print("Font file not found")
        return

    y_position = 30
        
    draw.text((70, y_position), f'{data["name"]} {data["middle_name"]} {data["last_name"]}', font=font_bold_big_size, fill=(0, 0, 0))
    y_position += 40
    draw.text((100, y_position), f'Age: {data["age"]}', font=font_regular, fill=(0, 0, 0))
    y_position += 23
    draw.text((100, y_position), f'Date of Birth: {data["dob"]}', font=font_regular, fill=(0, 0, 0))
    y_position += 23
    draw.text((100, y_position), f'Email: {data["email"]}', font=font_regular, fill=(0, 0, 0))
    y_position += 23
    draw.text((100, y_position), f'Citizenship: {data["citizenship"]}', font=font_regular, fill=(0, 0, 0))
    y_position += 23
    draw.text((100, y_position), f'City: {data["city"]}', font=font_regular, fill=(0, 0, 0))
    y_position += 30

    
    if data["socials"]:
        draw.text((80, y_position), 'Social Networks:', font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for social in data["socials"]:
            draw.text((100, y_position), f'{social["service"]}: {social["link"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23

    
    if data["projects"]:
        draw.text((80, y_position), 'Projects:', font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for project in data["projects"]:
            draw.text((100, y_position), f'Project Name: {project["name"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Time: {project["time"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Link: {project["link"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Description: {project["description"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23

    
    if data["experiences"]:
        draw.text((80, y_position), 'Experience:', font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for exp in data["experiences"]:
            draw.text((100, y_position), f'Company: {exp["company"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Job Title: {exp["title"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Period: {exp["period"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Description: {exp["description"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23

    
    if data["education"]:
        draw.text((80, y_position), 'Education:', font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for edu in data["education"]:
            draw.text((100, y_position), f'Institution: {edu["institution"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Period: {edu["period"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23
            draw.text((100, y_position), f'Field of Study: {edu["field"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23

    
    if data["languages"]:
        draw.text((80, y_position), 'Languages:', font=font_bold, fill=(0, 0, 0))
        y_position += 25
        for lang in data["languages"]:
            draw.text((100, y_position), f'{lang["language"]}: {lang["level"]}', font=font_regular, fill=(0, 0, 0))
            y_position += 23

    image.save(jpg_path)
    return jpg_path